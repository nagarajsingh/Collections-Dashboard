import re

import pdfplumber
from docx import Document


IMAGE_PATTERN = r"[a-zA-Z0-9\-\.]+\.azurecr\.io\/([a-zA-Z0-9\-]+):([a-zA-Z0-9_\.\-]+)"


def extract_services_from_text(text_content):
    results = []

    matches = re.findall(IMAGE_PATTERN, text_content)

    for service, tag in matches:
        results.append({
            "Select": True,
            "Service": service,
            "Image Tag": tag,
            "vendorImage": f"{service}:{tag}",
            "Extraction Method": "Rule-Based",
        })

    unique = []
    seen = set()

    for item in results:
        if item["vendorImage"] not in seen:
            seen.add(item["vendorImage"])
            unique.append(item)

    return unique


def extract_text_from_pdf(pdf_file):
    text_content = ""
    pdf_file.seek(0)

    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_content += page_text + "\n"

    return text_content.strip()


def extract_text_from_docx(docx_file):
    text_content = ""
    docx_file.seek(0)

    document = Document(docx_file)

    for paragraph in document.paragraphs:
        if paragraph.text:
            text_content += paragraph.text + "\n"

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            text_content += " | ".join(cells) + "\n"

    return text_content.strip()


def extract_text_from_document(uploaded_file):
    file_name = uploaded_file.name.lower()

    if file_name.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)

    if file_name.endswith(".docx"):
        return extract_text_from_docx(uploaded_file)

    raise ValueError("Unsupported file type. Upload PDF or DOCX.")


def extract_services_from_document(uploaded_file):
    text_content = extract_text_from_document(uploaded_file)
    return extract_services_from_text(text_content)
